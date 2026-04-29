from base64 import b64decode
from pathlib import Path
import re
from types import SimpleNamespace
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from openpyxl import Workbook
from openpyxl.drawing.image import Image as SpreadsheetImage
from PIL import Image as PilImage

from digester.images import MockImageAnalyzer
from digester.images.base import ImageAnalyzer
from digester.core.models import EmbeddedImage, ImageAnalysis
from digester.sources.registry import SourceRegistry


def _write_test_png(path: Path) -> None:
    path.write_bytes(
        b64decode("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+jK4QAAAAASUVORK5CYII=")
    )


def _write_docx_with_embedded_image(path: Path) -> None:
    image_path = path.with_suffix(".png")
    _write_test_png(image_path)
    document = Document()
    document.add_paragraph("Before image")
    image_paragraph = document.add_paragraph("Screenshot shows the confirmation dialog")
    image_paragraph.add_run().add_picture(str(image_path))
    document.add_paragraph("After image")
    document.save(path)


def _write_docx_with_table_image(path: Path) -> None:
    image_path = path.with_suffix(".png")
    _write_test_png(image_path)
    document = Document()
    document.add_paragraph("Before table")
    table = document.add_table(rows=1, cols=1)
    image_paragraph = table.cell(0, 0).paragraphs[0]
    image_paragraph.text = "Screenshot in a table cell"
    image_paragraph.add_run().add_picture(str(image_path))
    document.add_paragraph("After table")
    document.save(path)


def _write_docx_with_merged_table_image(path: Path) -> None:
    image_path = path.with_suffix(".png")
    _write_test_png(image_path)
    document = Document()
    table = document.add_table(rows=1, cols=2)
    merged_cell = table.cell(0, 0).merge(table.cell(0, 1))
    image_paragraph = merged_cell.paragraphs[0]
    image_paragraph.text = "Screenshot in a merged table cell"
    image_paragraph.add_run().add_picture(str(image_path))
    document.save(path)


def _write_docx_with_vml_image(path: Path) -> None:
    image_path = path.with_suffix(".png")
    _write_test_png(image_path)
    document = Document()
    document.add_picture(str(image_path))
    document.save(path)

    with ZipFile(path, "r") as source:
        document_xml = source.read("word/document.xml").decode("utf-8")
        match = re.search(r'r:embed="([^"]+)"', document_xml)
        assert match is not None
        rel_id = match.group(1)
        vml_object = (
            '<w:object><v:shape id="_x0000_i1025" type="#_x0000_t75">'
            '<v:imagedata r:id="{rel_id}"/></v:shape></w:object>'
        ).format(rel_id=rel_id)
        document_xml = re.sub(
            r"<w:drawing>.*?</w:drawing>",
            vml_object,
            document_xml,
            count=1,
            flags=re.DOTALL,
        )
        rewritten = path.with_suffix(".vml.docx")
        with ZipFile(rewritten, "w", ZIP_DEFLATED) as target:
            for item in source.infolist():
                data = source.read(item.filename)
                if item.filename == "word/document.xml":
                    data = document_xml.encode("utf-8")
                target.writestr(item, data)
    rewritten.replace(path)


def _write_docx_with_vml_vector_image(path: Path) -> None:
    _write_docx_with_vml_image(path)
    with ZipFile(path, "r") as source:
        content_types = source.read("[Content_Types].xml").decode("utf-8")
        content_types = content_types.replace(
            "</Types>",
            '<Default Extension="emf" ContentType="image/x-emf"/></Types>',
        )
        document_rels = source.read("word/_rels/document.xml.rels").decode("utf-8")
        document_rels = document_rels.replace("media/image1.png", "media/image1.emf")
        rewritten = path.with_suffix(".vector.docx")
        with ZipFile(rewritten, "w", ZIP_DEFLATED) as target:
            for item in source.infolist():
                if item.filename == "word/media/image1.png":
                    continue
                data = source.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    data = content_types.encode("utf-8")
                if item.filename == "word/_rels/document.xml.rels":
                    data = document_rels.encode("utf-8")
                target.writestr(item, data)
            target.writestr("word/media/image1.emf", source.read("word/media/image1.png"))
    rewritten.replace(path)


def _write_pdf_with_embedded_image(path: Path) -> None:
    image_path = path.with_suffix(".png")
    _write_test_png(image_path)
    with PilImage.open(image_path) as image:
        image.convert("RGB").save(path, "PDF")


def _write_spreadsheet_with_embedded_image(path: Path) -> None:
    image_path = path.with_suffix(".png")
    _write_test_png(image_path)
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Summary"
    sheet.append(["Item", "Value"])
    sheet.append(["alpha", 3])
    sheet.add_image(SpreadsheetImage(str(image_path)), "B2")
    workbook.save(path)


class RecordingReporter:
    def __init__(self) -> None:
        self.messages = []

    def update(self, message: str) -> None:
        self.messages.append(("update", message))

    def persist(self, message: str) -> None:
        self.messages.append(("persist", message))

    def verbose(self, message: str) -> None:
        self.messages.append(("verbose", message))

    def verbosity(self) -> int:
        return 0

    def clear(self) -> None:
        self.messages.append(("clear", ""))


class CapturingImageAnalyzer(ImageAnalyzer):
    def __init__(self) -> None:
        super().__init__()
        self.images = []

    def analyze(self, image: EmbeddedImage) -> ImageAnalysis:
        self.images.append(image)
        return ImageAnalysis(summary="Converted image was analyzed.", key_points=[])


def test_registry_loads_plain_text(tmp_path: Path) -> None:
    path = tmp_path / "notes.txt"
    path.write_text("alpha\n\nbeta", encoding="utf-8")

    documents = SourceRegistry().load_paths([path])

    assert len(documents) == 1
    assert documents[0].sections[0].content == "alpha\n\nbeta"


def test_registry_recursively_loads_supported_files_from_directories(tmp_path: Path) -> None:
    nested_dir = tmp_path / "docs" / "guides"
    nested_dir.mkdir(parents=True)
    root_path = tmp_path / "docs" / "overview.txt"
    nested_path = nested_dir / "setup.md"
    root_path.write_text("root document", encoding="utf-8")
    nested_path.write_text("nested document", encoding="utf-8")

    documents = SourceRegistry().load_paths([tmp_path / "docs"])

    loaded_paths = {document.path for document in documents}
    assert loaded_paths == {root_path, nested_path}


def test_registry_loads_docx(tmp_path: Path) -> None:
    path = tmp_path / "memo.docx"
    document = Document()
    document.add_paragraph("First paragraph")
    document.add_paragraph("Second paragraph")
    document.save(path)

    documents = SourceRegistry().load_paths([path])

    assert documents[0].sections[0].content == "First paragraph\n\nSecond paragraph"


def test_registry_detects_docx_embedded_images_without_analyzer(tmp_path: Path) -> None:
    path = tmp_path / "with-image.docx"
    _write_docx_with_embedded_image(path)

    documents = SourceRegistry().load_paths([path])

    assert len(documents[0].sections) == 1
    assert len(documents[0].embedded_images) == 1
    assert documents[0].embedded_images[0].caption == "Screenshot shows the confirmation dialog"
    assert documents[0].embedded_images[0].source_ref.locator == "embedded image 1 near paragraph 2"
    assert documents[0].embedded_images[0].mime_type == "image/png"
    assert len(documents[0].embedded_images[0].data) > 0
    assert documents[0].extraction_notes == [
        (
            "Detected 1 embedded image(s): embedded image 1 near paragraph 2: "
            "file=image1.png, mime=image/png, bytes=68."
        )
    ]
    assert documents[0].extraction_warnings == [
        "Detected 1 embedded image(s) but no image analyzer is configured; image content was skipped."
    ]


def test_registry_loads_docx_embedded_images_with_analyzer(tmp_path: Path) -> None:
    path = tmp_path / "with-image.docx"
    _write_docx_with_embedded_image(path)

    documents = SourceRegistry().load_paths([path], image_analyzer=MockImageAnalyzer(model="fake-vision"))

    assert len(documents[0].sections) == 2
    image_section = documents[0].sections[1]
    assert image_section.heading == "Embedded image 1"
    assert image_section.content_kind == "image-analysis"
    assert image_section.source_ref.locator == "embedded image 1 near paragraph 2"
    assert "Visual summary:" in image_section.content
    assert "Screenshot shows the confirmation dialog" in image_section.content
    assert documents[0].extraction_notes == [
        (
            "Detected 1 embedded image(s): embedded image 1 near paragraph 2: "
            "file=image1.png, mime=image/png, bytes=68."
        ),
        (
            "Analyzing embedded image 1: embedded image 1 near paragraph 2: "
            "file=image1.png, mime=image/png, bytes=68."
        ),
        "Analyzed embedded image 1 successfully.",
    ]
    assert documents[0].extraction_warnings == []


def test_registry_detects_docx_embedded_images_inside_tables(tmp_path: Path) -> None:
    path = tmp_path / "with-table-image.docx"
    _write_docx_with_table_image(path)

    documents = SourceRegistry().load_paths([path], image_analyzer=MockImageAnalyzer(model="fake-vision"))

    assert len(documents[0].embedded_images) == 1
    assert documents[0].embedded_images[0].caption == "Screenshot in a table cell"
    assert documents[0].embedded_images[0].source_ref.locator == "embedded image 1 near table paragraph 1"
    assert documents[0].embedded_images[0].context_text == "Before table\nAfter table"
    assert documents[0].embedded_images[0].mime_type == "image/png"
    assert len(documents[0].embedded_images[0].data) > 0
    assert len(documents[0].sections) == 2
    image_section = documents[0].sections[1]
    assert image_section.content_kind == "image-analysis"
    assert "Screenshot in a table cell" in image_section.content
    assert documents[0].extraction_warnings == []


def test_registry_deduplicates_docx_embedded_images_inside_merged_table_cells(tmp_path: Path) -> None:
    path = tmp_path / "with-merged-table-image.docx"
    _write_docx_with_merged_table_image(path)

    documents = SourceRegistry().load_paths([path], image_analyzer=MockImageAnalyzer(model="fake-vision"))

    assert len(documents[0].embedded_images) == 1
    assert documents[0].embedded_images[0].caption == "Screenshot in a merged table cell"
    assert len([section for section in documents[0].sections if section.content_kind == "image-analysis"]) == 1
    assert documents[0].extraction_warnings == []


def test_registry_detects_docx_vml_embedded_images(tmp_path: Path) -> None:
    path = tmp_path / "with-vml-image.docx"
    _write_docx_with_vml_image(path)

    documents = SourceRegistry().load_paths([path], image_analyzer=MockImageAnalyzer(model="fake-vision"))

    assert len(documents[0].embedded_images) == 1
    assert documents[0].embedded_images[0].source_ref.locator == "embedded image 1 near paragraph 1"
    assert documents[0].embedded_images[0].filename == "image1.png"
    assert documents[0].embedded_images[0].mime_type == "image/png"
    assert len(documents[0].embedded_images[0].data) > 0
    assert len(documents[0].sections) == 1
    assert documents[0].sections[0].content_kind == "image-analysis"
    assert "Visual summary:" in documents[0].sections[0].content
    assert documents[0].extraction_warnings == []


def test_registry_normalizes_vector_docx_images_before_analysis(monkeypatch, tmp_path: Path) -> None:
    path = tmp_path / "with-vector-image.docx"
    _write_docx_with_vml_vector_image(path)
    analyzer = CapturingImageAnalyzer()

    def fake_run(command, capture_output, text, timeout):
        output_path = Path(command[-1].split("=", 1)[1])
        _write_test_png(output_path)
        return SimpleNamespace(returncode=0, stdout="", stderr="")

    monkeypatch.setattr(
        "digester.sources.embedded_images._image_converter_command",
        lambda input_path, output_path: [
            "fake-inkscape",
            str(input_path),
            "--export-type=png",
            "--export-filename={output_path}".format(output_path=output_path),
        ],
    )
    monkeypatch.setattr("digester.sources.embedded_images.subprocess.run", fake_run)

    documents = SourceRegistry().load_paths([path], image_analyzer=analyzer)

    assert len(analyzer.images) == 1
    assert analyzer.images[0].filename == "image1.png"
    assert analyzer.images[0].mime_type == "image/png"
    assert analyzer.images[0].data.startswith(b"\x89PNG")
    assert any("Normalized embedded image for analysis" in note for note in documents[0].extraction_notes)
    assert documents[0].extraction_warnings == []


def test_registry_supports_legacy_inkscape_vector_conversion(monkeypatch, tmp_path: Path) -> None:
    path = tmp_path / "with-vector-image.docx"
    _write_docx_with_vml_vector_image(path)
    analyzer = CapturingImageAnalyzer()

    def fake_run(command, capture_output, text, timeout):
        if command == ["fake-inkscape", "--help"]:
            return SimpleNamespace(returncode=0, stdout="  -e, --export-png=FILENAME", stderr="")
        output_arg = [part for part in command if part.startswith("--export-png=")][0]
        output_path = Path(output_arg.split("=", 1)[1])
        _write_test_png(output_path)
        return SimpleNamespace(returncode=0, stdout="", stderr="")

    monkeypatch.setattr(
        "digester.sources.embedded_images.shutil.which",
        lambda name: "fake-inkscape" if name == "inkscape" else None,
    )
    monkeypatch.setattr("digester.sources.embedded_images.subprocess.run", fake_run)

    documents = SourceRegistry().load_paths([path], image_analyzer=analyzer)

    assert len(analyzer.images) == 1
    assert analyzer.images[0].mime_type == "image/png"
    assert documents[0].extraction_warnings == []


def test_registry_skips_vector_docx_images_when_normalization_fails(monkeypatch, tmp_path: Path) -> None:
    path = tmp_path / "with-vector-image.docx"
    _write_docx_with_vml_vector_image(path)
    analyzer = CapturingImageAnalyzer()

    monkeypatch.setattr(
        "digester.sources.embedded_images._image_converter_command",
        lambda input_path, output_path: ["fake-convert", str(input_path), str(output_path)],
    )
    monkeypatch.setattr(
        "digester.sources.embedded_images.subprocess.run",
        lambda command, capture_output, text, timeout: SimpleNamespace(
            returncode=1,
            stdout="",
            stderr="no decode delegate for this image format `EMF'",
        ),
    )

    documents = SourceRegistry().load_paths([path], image_analyzer=analyzer)

    assert analyzer.images == []
    assert documents[0].sections == []
    assert len(documents[0].embedded_images) == 1
    assert documents[0].extraction_warnings == [
        (
            "Skipped embedded image 1: Unable to convert embedded image 1 near paragraph 1: "
            "file=image1.emf, mime=image/x-emf, bytes=68 to PNG: "
            "no decode delegate for this image format `EMF'"
        )
    ]


def test_registry_logs_docx_embedded_image_metadata(tmp_path: Path) -> None:
    path = tmp_path / "with-image.docx"
    _write_docx_with_embedded_image(path)
    reporter = RecordingReporter()

    SourceRegistry().load_paths([path], progress_reporter=reporter)

    persisted = [message for kind, message in reporter.messages if kind == "persist"]
    assert any(
        message.startswith("Note for with-image.docx: Detected 1 embedded image(s): ")
        and "embedded image 1 near paragraph 2" in message
        and "file=image1.png, mime=image/png, bytes=68" in message
        for message in persisted
    )
    assert any(
        message == (
            "Warning for with-image.docx: "
            "Detected 1 embedded image(s) but no image analyzer is configured; image content was skipped."
        )
        for message in persisted
    )


class FailingImageAnalyzer(ImageAnalyzer):
    def analyze(self, image: EmbeddedImage):
        raise ValueError("vision request timed out")


def test_registry_keeps_docx_text_when_embedded_image_analysis_fails(tmp_path: Path) -> None:
    path = tmp_path / "with-image.docx"
    _write_docx_with_embedded_image(path)

    documents = SourceRegistry().load_paths([path], image_analyzer=FailingImageAnalyzer())

    assert len(documents[0].sections) == 1
    assert documents[0].sections[0].content == (
        "Before image\n\nScreenshot shows the confirmation dialog\n\nAfter image"
    )
    assert documents[0].embedded_images[0].source_ref.locator == "embedded image 1 near paragraph 2"
    assert documents[0].extraction_notes == [
        (
            "Detected 1 embedded image(s): embedded image 1 near paragraph 2: "
            "file=image1.png, mime=image/png, bytes=68."
        ),
        (
            "Analyzing embedded image 1: embedded image 1 near paragraph 2: "
            "file=image1.png, mime=image/png, bytes=68."
        ),
    ]
    assert documents[0].extraction_warnings == [
        "Failed to analyze embedded image 1: vision request timed out"
    ]


def test_registry_loads_spreadsheet(tmp_path: Path) -> None:
    path = tmp_path / "sheet.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Summary"
    sheet.append(["Item", "Value"])
    sheet.append(["alpha", 3])
    workbook.save(path)

    documents = SourceRegistry().load_paths([path])

    assert "Item | Value" in documents[0].sections[0].content
    assert "alpha | 3" in documents[0].sections[0].content


def test_registry_loads_pdf_embedded_images_with_analyzer(tmp_path: Path) -> None:
    path = tmp_path / "image-only.pdf"
    _write_pdf_with_embedded_image(path)

    documents = SourceRegistry().load_paths([path], image_analyzer=MockImageAnalyzer(model="fake-vision"))

    assert len(documents[0].sections) == 1
    assert documents[0].sections[0].content_kind == "image-analysis"
    assert documents[0].embedded_images[0].source_ref.locator == "embedded image 1 on page 1"
    assert documents[0].embedded_images[0].mime_type.startswith("image/")
    assert documents[0].extraction_notes[0].startswith("Detected 1 embedded image(s): embedded image 1 on page 1: ")
    assert "Analyzed embedded image 1 successfully." in documents[0].extraction_notes
    assert documents[0].extraction_warnings == ["Page 1 did not yield extractable text."]


def test_registry_loads_spreadsheet_embedded_images_with_analyzer(tmp_path: Path) -> None:
    path = tmp_path / "sheet-with-image.xlsx"
    _write_spreadsheet_with_embedded_image(path)

    documents = SourceRegistry().load_paths([path], image_analyzer=MockImageAnalyzer(model="fake-vision"))

    assert len(documents[0].sections) == 2
    assert documents[0].sections[0].heading == "Summary"
    assert documents[0].sections[1].content_kind == "image-analysis"
    assert documents[0].embedded_images[0].caption == "alpha | 3"
    assert documents[0].embedded_images[0].context_text == "Item | Value"
    assert documents[0].embedded_images[0].source_ref.locator == "embedded image 1 on sheet Summary near B2"
    assert "alpha | 3" in documents[0].sections[1].content
    assert "Item | Value" in documents[0].sections[1].content
    assert documents[0].extraction_warnings == []
