from __future__ import annotations

from dataclasses import replace
from pathlib import Path, PurePosixPath
import shutil
import subprocess
import tempfile
from typing import List, Optional, Tuple

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from ..core.models import DocumentSection, EmbeddedImage, ImageAnalysis, SourceDocument, SourceRef
from ..images.base import ImageAnalyzer
from .base import SourceAdapter

_BLIP_TAG = ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
_VML_IMAGE_TAG = ".//{urn:schemas-microsoft-com:vml}imagedata"
_EMBED_ATTR = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
_RELATIONSHIP_ID_ATTR = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
_MIME_TYPES_BY_SUFFIX = {
    ".emf": "image/x-emf",
    ".gif": "image/gif",
    ".jpeg": "image/jpeg",
    ".jpg": "image/jpeg",
    ".png": "image/png",
    ".tif": "image/tiff",
    ".tiff": "image/tiff",
    ".webp": "image/webp",
    ".wmf": "image/x-wmf",
}
_VECTOR_IMAGE_SUFFIXES = {".emf", ".wmf"}
_VECTOR_IMAGE_MIME_TYPES = {"image/emf", "image/x-emf", "image/wmf", "image/x-wmf"}


def _nearest_non_empty_paragraph_text(
    paragraphs: List[Paragraph],
    start_index: int,
    step: int,
) -> str:
    index = start_index + step
    while 0 <= index < len(paragraphs):
        text = paragraphs[index].text.strip()
        if text:
            return text
        index += step
    return ""


def _dedupe_non_empty(items: List[str]) -> List[str]:
    seen = set()
    result: List[str] = []
    for item in items:
        normalized = item.strip()
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        result.append(normalized)
    return result


def _mime_type_for_filename(filename: str) -> str:
    return _MIME_TYPES_BY_SUFFIX.get(PurePosixPath(filename).suffix.lower(), "application/octet-stream")


def _paragraphs_from_table(table: Table) -> List[Paragraph]:
    paragraphs: List[Paragraph] = []
    seen_cells = set()
    for row in table.rows:
        for cell in row.cells:
            if cell._tc in seen_cells:
                continue
            seen_cells.add(cell._tc)
            paragraphs.extend(cell.paragraphs)
            for nested_table in cell.tables:
                paragraphs.extend(_paragraphs_from_table(nested_table))
    return paragraphs


def _document_paragraphs_with_locations(document: DocxDocument) -> List[Tuple[str, int, Paragraph]]:
    located: List[Tuple[str, int, Paragraph]] = []
    paragraph_offset = 0
    table_paragraph_offset = 0
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            paragraph_offset += 1
            located.append(("paragraph", paragraph_offset, Paragraph(child, document)))
            continue
        if isinstance(child, CT_Tbl):
            for paragraph in _paragraphs_from_table(Table(child, document)):
                table_paragraph_offset += 1
                located.append(("table paragraph", table_paragraph_offset, paragraph))
    return located


def _locator_for_image(image_index: int, location_kind: str, location_offset: int) -> str:
    if location_kind == "table paragraph":
        return "embedded image {index} near table paragraph {paragraph}".format(
            index=image_index,
            paragraph=location_offset,
        )
    return "embedded image {index} near paragraph {paragraph}".format(
        index=image_index,
        paragraph=location_offset,
    )


def _image_relationship_ids(run) -> List[str]:
    rel_ids: List[str] = []
    for blip in run._element.findall(_BLIP_TAG):
        rel_id = str(blip.get(_EMBED_ATTR, "")).strip()
        if rel_id:
            rel_ids.append(rel_id)
    for image_data in run._element.findall(_VML_IMAGE_TAG):
        rel_id = str(image_data.get(_RELATIONSHIP_ID_ATTR, "")).strip()
        if rel_id:
            rel_ids.append(rel_id)
    return rel_ids


def _extract_embedded_images(
    document: DocxDocument,
    source_id: str,
    path: Path,
) -> List[EmbeddedImage]:
    located_paragraphs = _document_paragraphs_with_locations(document)
    paragraphs = [paragraph for _, _, paragraph in located_paragraphs]
    images: List[EmbeddedImage] = []
    image_index = 1
    for paragraph_offset, (location_kind, location_number, paragraph) in enumerate(located_paragraphs):
        caption = paragraph.text.strip()
        nearby_context = "\n".join(
            _dedupe_non_empty(
                [
                    _nearest_non_empty_paragraph_text(paragraphs, paragraph_offset, -1),
                    _nearest_non_empty_paragraph_text(paragraphs, paragraph_offset, 1),
                ]
            )
        )
        for run in paragraph.runs:
            for rel_id in _image_relationship_ids(run):
                relationship = document.part.rels.get(rel_id)
                image_part = document.part.related_parts.get(rel_id)
                if relationship is None or image_part is None:
                    continue
                filename = PurePosixPath(getattr(relationship, "target_ref", "")).name
                images.append(
                    EmbeddedImage(
                        image_id="{source_id}-image-{index}".format(
                            source_id=source_id,
                            index=image_index,
                        ),
                        source_ref=SourceRef(
                            source_id=source_id,
                            source_path=str(path),
                            locator=_locator_for_image(image_index, location_kind, location_number),
                        ),
                        filename=filename or "image-{index}".format(index=image_index),
                        mime_type=getattr(image_part, "content_type", "")
                        or _mime_type_for_filename(filename),
                        data=image_part.blob,
                        caption=caption,
                        context_text=nearby_context,
                    )
                )
                image_index += 1
    return images


def _render_image_analysis(image: EmbeddedImage, analysis: ImageAnalysis) -> str:
    lines = [
        "This section summarizes an embedded image from the source document.",
        "",
        "Visual summary: {summary}".format(summary=analysis.summary.strip()),
    ]
    if image.filename.strip():
        lines.extend(
            [
                "",
                "Image file: {filename}".format(filename=image.filename.strip()),
            ]
        )
    if image.caption.strip():
        lines.extend(
            [
                "",
                "Inline caption or nearby text: {caption}".format(caption=image.caption.strip()),
            ]
        )
    if image.context_text.strip():
        lines.extend(
            [
                "",
                "Nearby document context:",
                image.context_text.strip(),
            ]
        )
    if analysis.key_points:
        lines.extend(["", "Key visual details:"])
        lines.extend("- {point}".format(point=point) for point in analysis.key_points)
    return "\n".join(lines)


def _image_metadata_line(image: EmbeddedImage) -> str:
    return (
        "{locator}: file={filename}, mime={mime_type}, bytes={byte_count}"
    ).format(
        locator=image.source_ref.locator,
        filename=image.filename or "(unnamed)",
        mime_type=image.mime_type or "application/octet-stream",
        byte_count=len(image.data),
    )


def _image_requires_png_normalization(image: EmbeddedImage) -> bool:
    suffix = PurePosixPath(image.filename).suffix.lower()
    return suffix in _VECTOR_IMAGE_SUFFIXES or image.mime_type.lower() in _VECTOR_IMAGE_MIME_TYPES


def _image_converter_command(input_path: Path, output_path: Path) -> Optional[List[str]]:
    inkscape = shutil.which("inkscape")
    if inkscape:
        try:
            help_result = subprocess.run(
                [inkscape, "--help"],
                capture_output=True,
                text=True,
                timeout=10,
            )
            help_text = "{stdout}\n{stderr}".format(
                stdout=help_result.stdout,
                stderr=help_result.stderr,
            )
        except Exception:
            help_text = ""
        if "--export-png" in help_text:
            return [inkscape, "--without-gui", str(input_path), "--export-png={output_path}".format(
                output_path=output_path,
            )]
        return [
            inkscape,
            str(input_path),
            "--export-type=png",
            "--export-filename={output_path}".format(output_path=output_path),
        ]
    magick = shutil.which("magick")
    if magick:
        return [magick, str(input_path), str(output_path)]
    convert = shutil.which("convert")
    if convert:
        return [convert, str(input_path), str(output_path)]
    return None


def _normalize_image_for_analysis(image: EmbeddedImage) -> Tuple[Optional[EmbeddedImage], str]:
    if not _image_requires_png_normalization(image):
        return image, ""
    suffix = PurePosixPath(image.filename).suffix.lower() or ".img"
    try:
        with tempfile.TemporaryDirectory(prefix="bookworm-image-") as directory:
            input_path = Path(directory) / "input{suffix}".format(suffix=suffix)
            output_path = Path(directory) / "output.png"
            input_path.write_bytes(image.data)
            command = _image_converter_command(input_path=input_path, output_path=output_path)
            if command is None:
                return (
                    None,
                    (
                        "Image {details} uses a vector preview format that most vision APIs cannot decode, "
                        "and no Inkscape or ImageMagick converter was found."
                    ).format(details=_image_metadata_line(image)),
                )
            result = subprocess.run(
                command,
                capture_output=True,
                text=True,
                timeout=30,
            )
            if result.returncode != 0:
                detail = (result.stderr or result.stdout or "converter exited with a non-zero status").strip()
                return (
                    None,
                    "Unable to convert {details} to PNG: {detail}".format(
                        details=_image_metadata_line(image),
                        detail=detail,
                    ),
                )
            converted_data = output_path.read_bytes()
    except Exception as error:
        return (
            None,
            "Unable to convert {details} to PNG: {error}".format(
                details=_image_metadata_line(image),
                error=error,
            ),
        )
    converted_filename = "{stem}.png".format(stem=PurePosixPath(image.filename).stem or "image")
    converted_image = replace(
        image,
        filename=converted_filename,
        mime_type="image/png",
        data=converted_data,
    )
    return converted_image, "Normalized embedded image for analysis: {original} -> {converted}.".format(
        original=_image_metadata_line(image),
        converted=_image_metadata_line(converted_image),
    )


class DocxAdapter(SourceAdapter):
    supported_suffixes = (".docx",)
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def load(
        self,
        path: Path,
        image_analyzer: Optional[ImageAnalyzer] = None,
    ) -> SourceDocument:
        document = Document(str(path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        source_id = path.stem.replace(" ", "-").lower()
        sections: List[DocumentSection] = []
        if paragraphs:
            sections.append(
                DocumentSection(
                    heading=path.stem,
                    content="\n\n".join(paragraphs),
                    source_ref=SourceRef(
                        source_id=source_id,
                        source_path=str(path),
                        locator="document-body",
                    ),
                )
            )
        embedded_images = _extract_embedded_images(document=document, source_id=source_id, path=path)
        notes: List[str] = []
        warnings: List[str] = []
        if embedded_images:
            notes.append(
                "Detected {count} embedded image(s): {details}.".format(
                    count=len(embedded_images),
                    details="; ".join(_image_metadata_line(image) for image in embedded_images),
                )
            )
        if embedded_images and image_analyzer is None:
            warnings.append(
                "Detected {count} embedded image(s) but no image analyzer is configured; image content was skipped."
                .format(count=len(embedded_images))
            )
        if image_analyzer is not None:
            for index, image in enumerate(embedded_images, start=1):
                analysis_image, normalization_message = _normalize_image_for_analysis(image)
                if normalization_message and analysis_image is not None:
                    notes.append(normalization_message)
                if normalization_message and analysis_image is None:
                    warnings.append(
                        "Skipped embedded image {index}: {message}".format(
                            index=index,
                            message=normalization_message,
                        )
                    )
                    continue
                notes.append(
                    "Analyzing embedded image {index}: {details}.".format(
                        index=index,
                        details=_image_metadata_line(analysis_image),
                    )
                )
                try:
                    analysis = image_analyzer.analyze(analysis_image)
                    sections.append(
                        DocumentSection(
                            heading="Embedded image {index}".format(index=index),
                            content=_render_image_analysis(analysis_image, analysis),
                            source_ref=analysis_image.source_ref,
                            content_kind="image-analysis",
                        )
                    )
                    notes.append(
                        "Analyzed embedded image {index} successfully.".format(index=index)
                    )
                except Exception as error:
                    warnings.append(
                        "Failed to analyze embedded image {index}: {error}".format(
                            index=index,
                            error=error,
                        )
                    )
        return SourceDocument(
            source_id=source_id,
            path=path,
            media_type=self.media_type,
            title=path.stem,
            sections=sections,
            embedded_images=embedded_images,
            extraction_notes=notes,
            extraction_warnings=warnings,
        )
